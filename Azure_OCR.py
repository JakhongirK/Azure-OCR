import time
import os
from docx import Document
from docx.shared import Inches
import requests

subscription_key = "API"
endpoint = "https://ENDPOINT.cognitiveservices.azure.com/"
ocr_url = endpoint + "vision/v3.2/ocr"

folder_path = "C:\\Users\\"
output_folder = "C:\\Users\\"

if not os.path.exists(output_folder):
    os.makedirs(output_folder)

batch_size = 5 # Set the batch size
image_data_list = [] # List to hold the image data
docx_file_paths = [] # List to hold the output file paths

for root, dirs, files in os.walk(folder_path):
    for file_name in files:
        if file_name.endswith(".jpg"):
            image_path = os.path.join(root, file_name)
            try:
                with open(image_path, "rb") as image_file:
                    image_data = image_file.read()
                image_data_list.append(image_data)
                docx_file_paths.append(os.path.join(output_folder, file_name.split(".")[0] + ".docx"))
                if len(image_data_list) == batch_size:
                    headers = {'Ocp-Apim-Subscription-Key': subscription_key,
                               'Content-Type': 'application/octet-stream'}
                    params = {'language': 'unk', 'detectOrientation': 'true', 'mode': 'Handwritten'}
                    response = requests.post(ocr_url, headers=headers, params=params, data=b''.join(image_data_list))
                    response.raise_for_status()
                    analysis = response.json()
                    for i, image_data in enumerate(image_data_list):
                        text = ""
                        for region in analysis["recognitionResult"]["lines"][i]["regions"]:
                            for line in region["lines"]:
                                for word in line["words"]:
                                    text += word["text"] + " "
                                text += "\n"
                        docx_file_path = docx_file_paths[i]
                        document = Document()
                        document.add_picture(image_path, width=Inches(6))
                        document.add_paragraph(text)
                        document.save(docx_file_path)
                        time.sleep(1) # Add a delay of 1 second
                    image_data_list.clear()
                    docx_file_paths.clear()
            except Exception as e:
                print(f"Error processing file {file_name}: {e}")
                continue

if image_data_list: # Process any remaining images
    headers = {'Ocp-Apim-Subscription-Key': subscription_key,
               'Content-Type': 'application/octet-stream'}
    params = {'language': 'unk', 'detectOrientation': 'true', 'mode': 'Handwritten'}
    response = requests.post(ocr_url, headers=headers, params=params, data=b''.join(image_data_list))
    response.raise_for_status()
    analysis = response.json()
    for i, image_data in enumerate(image_data_list):
        text = ""
        for region in analysis["recognitionResult"]["lines"][i]["regions"]:
            for line in region["lines"]:
                for word in line["words"]:
                    text += word["text"] + " "
                text += "\n"
        docx_file_path = docx_file_paths[i]
        document = Document()
        document.add_picture(image_path, width=Inches(6))
        document.add_paragraph(text)
        document.save(docx_file_path)
        time.sleep(1) # Add a delay of 1 second
